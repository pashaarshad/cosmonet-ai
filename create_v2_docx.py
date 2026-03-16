import sys
import os

try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"C:\Users\Admin\.gemini\antigravity\brain\ff94b4d6-4acb-4e8b-8287-7b0470899156"

SECTIONS = [
    {
        "title": "Section 1 — Hero Banner & Navigation",
        "bg": "DARK NAVY",
        "img": "v2_hero_section_1773412524268.png",
    },
    {
        "title": "Section 2 — About Cosmonet AI + Key Metrics",
        "bg": "WHITE",
        "img": "v2_about_white_1773412541060.png",
    },
    {
        "title": "Section 3 — Our Services",
        "bg": "LIGHT GRAY",
        "img": "v2_services_white_1773412556796.png",
    },
    {
        "title": "Section 4 — Why Choose Cosmonet AI",
        "bg": "WHITE",
        "img": "v2_whychoose_white_1773412596591.png",
    },
    {
        "title": "Section 5 — Partnering with Cosmonet AI",
        "bg": "DARK NAVY",
        "img": "v2_partnership_dark_1773412616405.png",
    },
    {
        "title": "Section 6 — What We Offer",
        "bg": "WHITE",
        "img": "v2_offerings_white_1773412634172.png",
    },
    {
        "title": "Section 7 — Blogs & Insights",
        "bg": "LIGHT GRAY",
        "img": "v2_blogs_white_1773412666215.png",
    },
    {
        "title": "Section 8 — Careers",
        "bg": "WHITE",
        "img": "v2_careers_white_1773412684223.png",
    },
    {
        "title": "Section 9 — Contact Us",
        "bg": "DARK NAVY",
        "img": "v2_contact_dark_1773412704724.png",
    },
    {
        "title": "Section 10 — Footer",
        "bg": "DARKEST NAVY",
        "img": "v2_footer_dark_1773412752092.png",
    },
]

def create_v2_doc(docx_path):
    doc = docx.Document()

    # --- Cover Page ---
    for _ in range(6):
        doc.add_paragraph("")

    cover_title = doc.add_heading("COSMONET AI", level=0)
    cover_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Version 2 — UI/UX Design Mockups")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x9B, 0x9E)

    doc.add_paragraph("")

    tagline = doc.add_paragraph("Humanized • Mixed Backgrounds • White + Dark Navy\nProfessional Corporate Website Design")
    tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in tagline.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph("")

    date_line = doc.add_paragraph("March 2026")
    date_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in date_line.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xB0, 0xBE, 0xC5)

    doc.add_page_break()

    # --- Background Color Map Page ---
    bg_title = doc.add_heading("Background Color Map", level=1)

    bg_intro = doc.add_paragraph("This version alternates between white and dark backgrounds for a natural, human-designed feel:")
    bg_intro.runs[0].font.size = Pt(11)
    doc.add_paragraph("")

    # Create a simple color map table
    table = doc.add_table(rows=len(SECTIONS) + 1, cols=3)
    table.style = 'Light Shading Accent 1'

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Background"
    hdr[2].text = "Theme"

    for i, s in enumerate(SECTIONS):
        row = table.rows[i + 1].cells
        row[0].text = s["title"].split("—")[0].strip()
        row[1].text = s["bg"]
        if "WHITE" in s["bg"] or "LIGHT" in s["bg"]:
            row[2].text = "☀ Light"
        else:
            row[2].text = "🌙 Dark"

    doc.add_paragraph("")
    flow_note = doc.add_paragraph("Flow: Dark → White → Light Gray → White → Dark → White → Light Gray → White → Dark → Darkest")
    flow_note.runs[0].font.size = Pt(10)
    flow_note.runs[0].font.italic = True

    doc.add_page_break()

    # --- Section Pages (images only) ---
    for idx, section in enumerate(SECTIONS):
        # Section heading
        h = doc.add_heading(section["title"], level=1)

        # Background tag
        bg_tag = doc.add_paragraph(f"Background: {section['bg']}")
        bg_tag.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in bg_tag.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            if "WHITE" in section["bg"] or "LIGHT" in section["bg"]:
                run.font.color.rgb = RGBColor(0x1A, 0x9B, 0x9E)  # teal
            else:
                run.font.color.rgb = RGBColor(0xE8, 0x62, 0x2B)  # orange

        # Insert image
        img_path = os.path.join(BASE, section["img"])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.2))
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph(f"[Image not found: {img_path}]")

        # Page break after each section except the last
        if idx < len(SECTIONS) - 1:
            doc.add_page_break()

    # --- Final page ---
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph("")

    end_title = doc.add_heading("Ready for Development", level=1)
    end_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    end_note = doc.add_paragraph("All section mockups above are finalized for the Cosmonet AI homepage.\nProportional layout, mixed light/dark backgrounds, humanized design.")
    end_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in end_note.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(docx_path)
    print(f"V2 Design Document saved at: {docx_path}")

if __name__ == "__main__":
    out = r"d:\cosmonet-ai\Cosmonet_AI_UIUX_V2_Design.docx"
    create_v2_doc(out)
