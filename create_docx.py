import sys
import os

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def create_doc(md_path, docx_path):
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('COSMONET AI \nComplete UI/UX Design Document', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    # Regex to extract image paths: ![alt text](file_path)
    img_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check for image first
        img_match = img_pattern.search(line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Remove any trailing ' ' or '`' just in case
            img_path = img_path.strip("'` ")
            
            try:
                # Add the image
                if os.path.exists(img_path):
                    # Most images should fit within the page width (typically ~6 inches margins considered)
                    # We'll set width to 6 inches, docx will scale height proportionally
                    doc.add_picture(img_path, width=Inches(6.0))
                    
                    # Add caption
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption.style = 'Caption'
                else:
                    doc.add_paragraph(f"[Image not found at path: {img_path}]")
            except Exception as e:
                doc.add_paragraph(f"[Error inserting image: {str(e)}]")
            
            # Since the line was an image, we can skip text processing for it,
            # unless there's more text. We will assume the image is on its own line.
            continue
            
        # Clean up some markdown specific syntax for word
        line = line.replace('**', '')
        line = line.replace('`', '')
        line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line) # Remove link URLs and keep text
        
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            if "COSMONET AI" in line: continue # Skip main title
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        elif line.startswith('|') and '---' in line:
            continue # skip table separators
        elif line.startswith('|'):
            # Basic handling for tables by adding text
            doc.add_paragraph(line) 
        elif line == '<!-- slide -->' or line == '````carousel' or line == '````':
            continue # Skip carousel meta tags
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Document saved successfully with images at: {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\Admin\.gemini\antigravity\brain\ff94b4d6-4acb-4e8b-8287-7b0470899156\cosmonet_ai_complete_uiux_doc.md"
    docx_file = r"d:\cosmonet-ai\Cosmonet_AI_UIUX_Plan_With_Images.docx"
    create_doc(md_file, docx_file)
